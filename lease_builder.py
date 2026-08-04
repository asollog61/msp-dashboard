"""Lease Builder helpers for the MSP Property Dashboard."""

from __future__ import annotations

import copy
from io import BytesIO
import json
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = BASE_DIR / "data" / "Lease Builder"
PUBLISHED_DIR = TEMPLATE_DIR / "Published"
# Hand-authored masters were moved into their own subfolder so the data
# directory stops mixing source documents with generated ones. Both places are
# scanned, because older checkouts still keep them at the top level.
MASTERS_DIR = TEMPLATE_DIR / "Templates"
PUBLISHED_PREFIX = "📦 "
CLAUSE_LIBRARY_FILE = TEMPLATE_DIR / "lease_clause_library.json"
SECTION_RE = re.compile(
    r"^\s*Section\s+(\d+(?:\.\d+)?)(?:\.(?:\s+|$)|\s+)(.*)$",
    re.IGNORECASE,
)

CATEGORY_RANGES = [
    (1, 9, "Lease Economics & Delivery"),
    (10, 17, "Premises Operations"),
    (18, 26, "Transfers, Casualty & Property Rights"),
    (27, 40, "Defaults, Remedies & Enforcement"),
    (41, 46, "Insurance, Indemnity & Environmental"),
    (47, 56, "General & Boilerplate"),
]


def discover_templates() -> list[dict[str, str]]:
    """Return deployable DOCX templates; future templates are auto-discovered.

    Hand-authored masters live in data/Lease Builder. Templates published from
    the dashboard live in the Published subfolder and are labelled so the two are
    never confused in the picker.
    """
    templates = []
    if not TEMPLATE_DIR.exists():
        return templates

    def add_from(directory: Path, prefix: str = "") -> None:
        if not directory.exists():
            return
        for path in sorted(directory.glob("*.docx")):
            if path.name.startswith("~$") or " TEST" in path.stem.upper():
                continue
            label = re.sub(r"^\d{4}_\d{2}_\d{2}\s+", "", path.stem)
            label = re.sub(r"\s*\.v(\d+)\s*", r" — v\1 ", label, flags=re.IGNORECASE)
            templates.append({"label": (prefix + label.strip()).strip(), "path": str(path)})

    add_from(TEMPLATE_DIR)
    add_from(MASTERS_DIR)
    add_from(PUBLISHED_DIR, prefix=PUBLISHED_PREFIX)
    return templates


def publish_template_docx(
    template_path: str | Path,
    name: str,
    section_choices: dict[str, dict[str, Any]],
    bookmark_values: dict[str, str] | None = None,
    rent_schedules: dict[str, Any] | None = None,
    clean_drafting_notes: bool = True,
    key_provision_rows: list[dict[str, Any]] | None = None,
) -> Path:
    """Write a resolved template to Published/ and return the new file path.

    The result is a normal Word file, so it carries its own formatting and can be
    hand-edited, and it is a valid base template: every section stays included
    and every key-provision bookmark is preserved.
    """
    safe_name = re.sub(r"[^A-Za-z0-9 ._-]+", "_", str(name)).strip(" _.") or "Published Template"
    PUBLISHED_DIR.mkdir(parents=True, exist_ok=True)
    destination = PUBLISHED_DIR / f"{safe_name}.docx"

    payload = build_word_document(
        template_path,
        section_choices=section_choices,
        bookmark_values=bookmark_values or {},
        included_bookmarks=None,
        linked_provisions=[],
        custom_provisions=[],
        rent_schedules=rent_schedules,
        additional_choices=[],
        clean_drafting_notes=clean_drafting_notes,
        document_title=safe_name,
        preserve_template_structure=True,
        key_provision_rows=key_provision_rows,
    )
    destination.write_bytes(payload)

    # A published file that cannot be parsed back is worse than no file, so it is
    # validated before being handed to the caller.
    check = inspect_template(destination)
    if not check["sections"]:
        destination.unlink(missing_ok=True)
        raise ValueError(
            "The published file could not be read back as a template "
            "(no numbered sections were found). Nothing was saved."
        )
    return destination


def load_clause_library() -> dict[str, Any]:
    try:
        with CLAUSE_LIBRARY_FILE.open("r", encoding="utf-8") as handle:
            data = json.load(handle)
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {"sections": {}, "additional_provisions": {}}


def category_for(number: str) -> str:
    try:
        integer = int(float(number))
    except ValueError:
        return "Other"
    for start, end, category in CATEGORY_RANGES:
        if start <= integer <= end:
            return category
    return "Other"


def _section_title(remainder: str, number: str) -> str:
    remainder = " ".join(remainder.split()).strip()
    if not remainder:
        return f"Section {number}"
    title = remainder.split(".", 1)[0].strip(" .\t")
    return title[:120] or f"Section {number}"


def scan_sections(document: Document) -> list[dict[str, Any]]:
    """Locate the numbered lease body, excluding drafting notes and exhibits."""
    paragraphs = document.paragraphs
    matches: list[tuple[int, str, str]] = []
    for index, paragraph in enumerate(paragraphs):
        match = SECTION_RE.match(paragraph.text.strip())
        if match:
            number, remainder = match.groups()
            matches.append((index, number, remainder))

    # The body begins at "Section 1". Anything before it is preamble or drafting
    # notes, which may themselves cite a section (the master lease mentions
    # Section 12 in its notes). Position is not a reliable cue: a generated
    # template has its notes stripped, so Section 1 sits near the top. When more
    # than one candidate exists (e.g. a table of contents), take the one that
    # yields the most sections, which is always the real body.
    def collect(from_position: int) -> list[tuple[int, str, str]]:
        picked: list[tuple[int, str, str]] = []
        for index, number, remainder in matches[from_position:]:
            picked.append((index, number, remainder))
            if number == "56":
                break
        return picked

    candidates = [position for position, item in enumerate(matches) if item[1] == "1"]
    starts = max((collect(position) for position in candidates), key=len, default=[])

    sections = []
    for position, (start, number, remainder) in enumerate(starts):
        if position + 1 < len(starts):
            end = starts[position + 1][0]
        else:
            end = next(
                (i for i in range(start + 1, len(paragraphs))
                 if paragraphs[i].text.strip().upper() == "SIGNATURES"),
                len(paragraphs),
            )
        clause_text = "\n\n".join(
            paragraph.text.strip()
            for paragraph in paragraphs[start:end]
            if paragraph.text.strip()
        )
        sections.append({
            "number": number,
            "title": _section_title(remainder, number),
            "category": category_for(number),
            "start": start,
            "end": end,
            "text": clause_text,
            "paragraph_count": end - start,
        })
    return sections


BOOKMARK_LABELS = {
    "Tx_BuildingAddress": "Property",
    "Tx_LeaseEffeftiveDt": "Lease Execution Date",
    "Tx_Landlord": "Landlord",
    "Tx_Tenant": "Tenant",
    "Tx_Premises": "Premises",
    "Tx_Sqft": "Premises Sqft",
    "Tx_TenantShareProperty": "Tenant Share of Property",
    "Tx_TenantShareFloor": "Tenant Share of Floor",
    "Tx_TenantSharePremisis": "Tenant Share of Premises",
    "Tx_PermittedUse": "Permitted Use",
    "Tx_LandlordContact": "Landlord Contact",
    "Tx_TenantContact": "Tenant Contact",
    "Tx_LeaseType": "Lease Type",
    "Tx_LeaseCommenceDt": "Lease Commencement Date",
    "Tx_DeliveryCondition": "Lease Delivery Condition",
    "Tx_RentCommDt": "Rent Commencement Date",
    "Tx_PermitPer": "Permit Period",
    "Tx_LeaseExpDt": "Lease Expiration Date",
    "Tx_LeaseTerm": "Lease Term",
    "Tx_AdditionalRent": "Additional Rent",
    "Tx_OptionPeriod": "Option Period(s)",
    "Tx_OptiontoCancel": "Option to Cancel",
    "Tx_Utilities": "Utilities and Services",
    "Tx_SecurityDeposit": "Security Deposit",
    "Tx_Brokers": "Broker(s)",
    "Tx_LandlordContFitUp": "Landlord Fit-Up Contribution",
}


def _bookmark_ranges(document: Document) -> dict[str, tuple[Any, Any, list[Any]]]:
    elements = list(document.element.body.iter())
    end_positions = {
        element.get(qn("w:id")): index
        for index, element in enumerate(elements)
        if element.tag == qn("w:bookmarkEnd")
    }
    ranges = {}
    for index, element in enumerate(elements):
        if element.tag != qn("w:bookmarkStart"):
            continue
        name = element.get(qn("w:name"))
        end_index = end_positions.get(element.get(qn("w:id")))
        if not name or end_index is None:
            continue
        ranges[name] = (element, elements[end_index], elements[index + 1:end_index])
    return ranges


def inspect_template(template_path: str | Path) -> dict[str, Any]:
    document = Document(str(template_path))
    bookmarks = []
    for name, (_, _, elements) in _bookmark_ranges(document).items():
        if name not in BOOKMARK_LABELS:
            continue
        value = "".join(
            element.text or "" for element in elements if element.tag == qn("w:t")
        ).strip()
        bookmarks.append({"bookmark": name, "field": BOOKMARK_LABELS[name], "value": value})
    bookmarks.sort(key=lambda item: list(BOOKMARK_LABELS).index(item["bookmark"]))

    # Some REF fields point at bookmarks that were never surfaced as key
    # provisions. They are reported so the caller can create them instead of
    # letting the cross-reference disappear.
    known = {item["bookmark"] for item in bookmarks}
    ref_targets = collect_ref_targets(document)
    orphan_refs = [
        {"bookmark": target, "field": BOOKMARK_LABELS.get(target, humanize_bookmark(target))}
        for target in ref_targets if target not in known
    ]

    # Clause text is read after the REF fields become KP: tokens, so the editor
    # shows KP:Landlord where Word would show "Error! Reference source not found."
    names_by_id = {item["bookmark"]: item["field"] for item in bookmarks}
    names_by_id.update({item["bookmark"]: item["field"] for item in orphan_refs})
    convert_ref_fields_to_tokens(document, names_by_id)

    return {
        "sections": scan_sections(document),
        "bookmarks": bookmarks,
        "orphan_refs": orphan_refs,
    }


def _replace_bookmark_text(document: Document, name: str, value: str) -> bool:
    target = _bookmark_ranges(document).get(name)
    if not target:
        return False
    start, _, elements = target
    text_elements = [element for element in elements if element.tag == qn("w:t")]
    if text_elements:
        text_elements[0].text = value
        if value.startswith(" ") or value.endswith(" "):
            text_elements[0].set(qn("xml:space"), "preserve")
        for element in text_elements[1:]:
            element.text = ""
    else:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        start.addnext(run)
    return True


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _ancestor(element: Any, tag: str) -> Any | None:
    current = element
    while current is not None:
        if current.tag == tag:
            return current
        current = current.getparent()
    return None


def _relocate_empty_bookmark(document: Document, start: Any, end: Any) -> None:
    """Keep an excluded bookmark valid while removing its visible source row."""
    for element in (start, end):
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_properties.append(OxmlElement("w:vanish"))
    run.append(run_properties)
    text = OxmlElement("w:t")
    text.text = ""
    run.append(text)
    paragraph.extend([start, run, end])

    body = document.element.body
    section_properties = body.find(qn("w:sectPr"))
    if section_properties is not None:
        section_properties.addprevious(paragraph)
    else:
        body.append(paragraph)


def _apply_key_provision_inclusions(document: Document, included_bookmarks: set[str]) -> None:
    """Remove fully excluded KPS rows while preserving blank REF sources."""
    ranges = _bookmark_ranges(document)
    row_groups: dict[int, dict[str, Any]] = {}
    for name in BOOKMARK_LABELS:
        target = ranges.get(name)
        if not target:
            continue
        row = _ancestor(target[0], qn("w:tr"))
        if row is None:
            if name not in included_bookmarks:
                _replace_bookmark_text(document, name, "")
            continue
        group = row_groups.setdefault(id(row), {"row": row, "names": []})
        group["names"].append(name)

    for group in row_groups.values():
        names = group["names"]
        excluded = [name for name in names if name not in included_bookmarks]
        if not excluded:
            continue
        if len(excluded) == len(names):
            current_ranges = _bookmark_ranges(document)
            for name in names:
                target = current_ranges.get(name)
                if target:
                    _relocate_empty_bookmark(document, target[0], target[1])
            row = group["row"]
            parent = row.getparent()
            if parent is not None:
                parent.remove(row)
        else:
            for name in excluded:
                _replace_bookmark_text(document, name, "")


# ---------------------------------------------------------------------------
# App-owned key provisions.
#
# The master lease carries hand-placed Tx_* bookmarks that were used both to
# discover the provision list and to inject values. They are brittle: a blank
# value could delete the row and take the bookmark with it, and provisions could
# never be added or renamed. The provision list now lives in the app, and the
# summary table is rebuilt from it on every build.
#
# Linking no longer copies a provision's text into the clause. Instead each
# included section gets a generated anchor bookmark and the summary row carries
# a hyperlink to it. Those anchors are disposable and rewritten on every build,
# unlike the hand-placed bookmarks they replace.
# ---------------------------------------------------------------------------

ANCHOR_PREFIX = "_MSP_Sec_"

# ---------------------------------------------------------------------------
# KP: tokens.
#
# The master lease cross-referenced key provisions with Word REF fields aimed at
# the hand-placed bookmarks. With those gone the fields render as "Error!
# Reference source not found." Clause text now cites a provision by name instead,
# as KP:Landlord, and the value is substituted at build time.
#
# There is no closing delimiter, so the end of a name is found by matching
# against the real provision list longest-name-first. That way KP:Lease
# Execution Date resolves to that provision rather than stopping at "Lease".
# ---------------------------------------------------------------------------

# The closing bracket delimits the name, so multi-word provisions need no
# guessing and a misspelled name is reported exactly as typed.
KP_TOKEN_RE = re.compile(r"\[\s*KPS?\s*:\s*([^\]]+?)\s*\]", re.IGNORECASE)
KP_ANCHOR_PREFIX = "_MSP_KP_"
# Green so a substituted cross-reference is obvious against the black body text.
KP_LINK_COLOR = "1F7A33"


def kp_token(name: str) -> str:
    return f"[KP:{name}]"


# A caret starts a new line inside a key-provision value, so an address can be
# written as "123 Main St^Westfield, NJ^07090". The grid's cell editor commits on
# Enter and cannot hold a real newline, so a visible marker is needed there; an
# Alt+Enter newline typed in Excel means the same thing and is folded into it.
KP_LINE_BREAK = "^"


def kp_value_lines(value: Any) -> list[str]:
    """Split a key-provision value into its display lines."""
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    return [line.strip() for line in re.split(r"\^|\n", text)]


def normalize_kp_value(value: Any) -> str:
    """Canonical single-line storage form: real newlines become carets."""
    return KP_LINE_BREAK.join(kp_value_lines(value))


def kp_value_plain(value: Any) -> str:
    """Value with real newlines, for Excel cells and previews."""
    return "\n".join(kp_value_lines(value))


def _kp_anchor_name(field: str) -> str:
    return KP_ANCHOR_PREFIX + re.sub(r"[^A-Za-z0-9_]", "_", str(field))[:24]


def humanize_bookmark(bookmark: str) -> str:
    """Tx_LandlordWork -> 'Landlord Work', for provisions that had no label."""
    stem = re.sub(r"^Tx_", "", str(bookmark))
    spaced = re.sub(r"(?<=[a-z0-9])(?=[A-Z])", " ", stem)
    spaced = re.sub(r"[_\-]+", " ", spaced)
    return re.sub(r"\s+", " ", spaced).strip() or str(bookmark)


def normalize_kp_tokens(text: str, names: Any) -> str:
    """Upgrade bare KP:Name citations to the bracketed [KP:Name] form.

    Clause text drafted before brackets existed, or typed without them, is
    repaired by matching the known provision names longest-first so that
    KP:Tenant Share of Property does not stop at KP:Tenant.
    """
    source = str(text or "")
    if ":" not in source:
        return source
    ordered = sorted(
        {str(name).strip() for name in names if str(name).strip()}, key=len, reverse=True
    )
    if not ordered:
        return source
    alternation = "|".join(re.escape(name) for name in ordered)
    # (?<!\[) leaves already-bracketed tokens untouched.
    pattern = re.compile(r"(?<!\[)\bKPS?:\s*(" + alternation + r")", re.IGNORECASE)
    lookup = {name.lower(): name for name in ordered}
    return pattern.sub(
        lambda match: kp_token(lookup.get(match.group(1).strip().lower(), match.group(1).strip())),
        source,
    )


def _kp_segments(text: str, values: dict[str, str], lookup: dict[str, str]) -> tuple[list, set, set]:
    """Split text into plain and cross-reference pieces.

    Each piece is (kind, text, provision_name). An unrecognised token is left
    visible so it cannot slip into a signed lease unnoticed.
    """
    segments: list[tuple[str, str, str | None]] = []
    referenced: set[str] = set()
    unresolved: set[str] = set()
    position = 0
    source = str(text or "")
    for match in KP_TOKEN_RE.finditer(source):
        if position < match.start():
            segments.append(("text", source[position:match.start()], None))
        raw = match.group(1).strip()
        actual = lookup.get(raw.lower())
        if actual is None:
            unresolved.add(raw)
            segments.append(("text", match.group(0), None))
        else:
            referenced.add(actual)
            # An unused provision resolves to nothing; the caller reports it.
            segments.append(("link", str(values.get(actual, "") or ""), actual))
        position = match.end()
    if position < len(source):
        segments.append(("text", source[position:], None))
    return segments, referenced, unresolved


def resolve_kp_text(text: str, values: dict[str, str]) -> tuple[str, set[str], set[str]]:
    """Plain-text substitution of [KP:Name] tokens, used for previews."""
    lookup = {str(name).strip().lower(): name for name in values}
    segments, referenced, unresolved = _kp_segments(text, values, lookup)
    return "".join(piece for _, piece, _ in segments), referenced, unresolved


def find_kp_references(text: str, names: Any) -> set[str]:
    """Provision names cited by [KP:Name] tokens in a piece of clause text."""
    lookup = {str(name).strip().lower(): str(name) for name in names}
    found = set()
    for match in KP_TOKEN_RE.finditer(str(text or "")):
        actual = lookup.get(match.group(1).strip().lower())
        if actual:
            found.add(actual)
    return found


def collect_ref_targets(document: Document) -> list[str]:
    """Bookmark names cited by Word REF fields, in document order."""
    targets = []
    for element in document.element.body.iter(qn("w:instrText")):
        match = re.search(r"\bREF\s+(\S+)", element.text or "", re.IGNORECASE)
        if match and match.group(1) not in targets:
            targets.append(match.group(1))
    return targets


def _convert_ref_fields_in_paragraph(paragraph: Paragraph, names_by_id: dict[str, str]) -> int:
    """Replace each REF field in a paragraph with a KP:<name> text run."""
    body = paragraph._p
    children = list(body)

    def field_type(element: Any) -> str | None:
        if element.tag != qn("w:r"):
            return None
        marker = element.find(qn("w:fldChar"))
        return marker.get(qn("w:fldCharType")) if marker is not None else None

    spans: list[tuple[int, int, str]] = []
    depth = 0
    start_index = 0
    instructions: list[str] = []
    for index, element in enumerate(children):
        kind = field_type(element)
        if kind == "begin":
            if depth == 0:
                start_index, instructions = index, []
            depth += 1
        elif depth > 0:
            if element.tag == qn("w:r"):
                instruction = element.find(qn("w:instrText"))
                if instruction is not None and instruction.text:
                    instructions.append(instruction.text)
            if kind == "end":
                depth -= 1
                if depth == 0:
                    spans.append((start_index, index, "".join(instructions)))

    converted = 0
    for start_index, end_index, instruction in reversed(spans):
        match = re.search(r"\bREF\s+(\S+)", instruction, re.IGNORECASE)
        if not match:
            continue
        name = names_by_id.get(match.group(1))
        if not name:
            continue
        group = children[start_index:end_index + 1]
        # Model the new run on the field's result run so the text keeps its look.
        prototype = next(
            (element for element in group
             if element.tag == qn("w:r") and element.find(qn("w:t")) is not None),
            None,
        )
        run = copy.deepcopy(prototype) if prototype is not None else OxmlElement("w:r")
        for tag in (qn("w:t"), qn("w:fldChar"), qn("w:instrText")):
            for stale in list(run.findall(tag)):
                run.remove(stale)
        text_element = OxmlElement("w:t")
        text_element.text = kp_token(name)
        text_element.set(qn("xml:space"), "preserve")
        run.append(text_element)
        body.insert(start_index, run)
        for element in group:
            if element.getparent() is not None:
                body.remove(element)
        converted += 1
    return converted


def _iter_all_paragraphs(document: Document) -> Any:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def convert_ref_fields_to_tokens(document: Document, names_by_id: dict[str, str]) -> int:
    """Rewrite every resolvable REF field in the document as a KP: token."""
    return sum(
        _convert_ref_fields_in_paragraph(paragraph, names_by_id)
        for paragraph in _iter_all_paragraphs(document)
    )


def _clone_run(prototype: Any, text: str, split_lines: bool = False) -> Any:
    """A run carrying the paragraph's existing character formatting.

    split_lines applies only to substituted provision values; ordinary clause
    prose is never split, so a stray caret in the lease text is left alone.
    """
    if prototype is not None:
        run = copy.deepcopy(prototype)
        for tag in (qn("w:t"), qn("w:fldChar"), qn("w:instrText"), qn("w:br")):
            for stale in list(run.findall(tag)):
                run.remove(stale)
    else:
        run = OxmlElement("w:r")
    _fill_run_lines(run, kp_value_lines(text) if split_lines else [str(text)])
    return run


def _green_link_run(prototype: Any, text: str) -> Any:
    """A cross-reference run: green and underlined so it reads as a link."""
    run = _clone_run(prototype, text, split_lines=True)
    properties = run.find(qn("w:rPr"))
    if properties is None:
        properties = OxmlElement("w:rPr")
        run.insert(0, properties)
    for tag in (qn("w:color"), qn("w:u")):
        for stale in list(properties.findall(tag)):
            properties.remove(stale)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), KP_LINK_COLOR)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    return run


def substitute_kp_tokens(
    document: Document,
    values: dict[str, str],
    kp_anchors: dict[str, str] | None = None,
) -> tuple[set[str], set[str]]:
    """Replace [KP:Name] tokens throughout the document.

    A resolved value becomes a green underlined hyperlink pointing back up to
    that provision's row in the Key Provisions summary. Returns the provision
    names referenced and any token names that matched nothing.
    """
    kp_anchors = kp_anchors or {}
    lookup = {str(name).strip().lower(): name for name in values}
    referenced: set[str] = set()
    unresolved: set[str] = set()

    for paragraph in _iter_all_paragraphs(document):
        original = paragraph.text
        if "[" not in original or ":" not in original:
            continue
        segments, found, missing = _kp_segments(original, values, lookup)
        referenced |= found
        unresolved |= missing
        if not any(kind == "link" for kind, _, _ in segments) and not missing:
            continue

        runs = paragraph.runs
        prototype = runs[0]._element if runs else None
        body = paragraph._p
        for child in list(body):
            if child.tag in (qn("w:r"), qn("w:hyperlink")):
                body.remove(child)

        for kind, piece, name in segments:
            if not piece:
                continue  # An excluded provision contributes nothing.
            if kind == "text":
                body.append(_clone_run(prototype, piece))
                continue
            anchor = kp_anchors.get(name)
            run = _green_link_run(prototype, piece)
            if anchor:
                link = OxmlElement("w:hyperlink")
                link.set(qn("w:anchor"), anchor)
                link.append(run)
                body.append(link)
            else:
                body.append(run)

    return referenced, unresolved


def _anchor_name(section_number: str) -> str:
    """Word bookmark names allow no dots and must start with a letter or underscore."""
    return ANCHOR_PREFIX + re.sub(r"[^A-Za-z0-9_]", "_", str(section_number))[:24]


def _add_section_anchors(document: Document) -> dict[str, str]:
    """Bookmark the first paragraph of every surviving section; return number -> anchor."""
    anchors: dict[str, str] = {}
    paragraphs = document.paragraphs
    bookmark_id = 9000
    for section in scan_sections(document):
        number = str(section["number"])
        start_index = section["start"]
        if start_index >= len(paragraphs):
            continue
        anchor = _anchor_name(number)
        paragraph = paragraphs[start_index]
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), anchor)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        # w:pPr must stay the first child of w:p, so the anchor goes after it.
        properties = paragraph._p.find(qn("w:pPr"))
        paragraph._p.insert(1 if properties is not None else 0, start)
        paragraph._p.append(end)
        anchors[number] = anchor
        bookmark_id += 1
    return anchors


def _fill_run_lines(run_element: Any, lines: list[str]) -> None:
    """Write one or more lines into a run, separated by soft line breaks."""
    for tag in (qn("w:t"), qn("w:br")):
        for stale in list(run_element.findall(tag)):
            run_element.remove(stale)
    for index, line in enumerate(lines or [""]):
        if index:
            run_element.append(OxmlElement("w:br"))
        text_element = OxmlElement("w:t")
        text_element.text = line
        text_element.set(qn("xml:space"), "preserve")
        run_element.append(text_element)


def _set_cell_text(cell: Any, text: str) -> Paragraph:
    """Replace a cell's text, honouring caret line breaks, keeping run formatting."""
    paragraphs = cell.paragraphs
    first = paragraphs[0]
    for paragraph in paragraphs[1:]:
        _delete_paragraph(paragraph)
    runs = first.runs
    if runs:
        target = runs[0]._element
        for run in runs[1:]:
            parent = run._element.getparent()
            if parent is not None:
                parent.remove(run._element)
    else:
        target = first.add_run("")._element
    _fill_run_lines(target, kp_value_lines(text))
    return first


def _append_hyperlink(paragraph: Paragraph, anchor: str, text: str) -> None:
    """Append an internal hyperlink that jumps to a bookmark in this document."""
    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    text_element.set(qn("xml:space"), "preserve")
    run.append(text_element)
    link.append(run)
    paragraph._p.append(link)


def _rebuild_key_provision_table(
    document: Document,
    provisions: list[dict[str, Any]],
    anchors: dict[str, str],
) -> dict[str, str]:
    """Replace every row of the summary table with the app's provision list.

    An existing row is cloned as the prototype so column widths, shading, fonts
    and the merge across the two value columns are carried over exactly.
    """
    if not document.tables:
        return {}
    table = document.tables[0]

    prototype = None
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2 and cells[0].text.strip() and cells[1].text.strip():
            # Prefer a row whose value cell spans the remaining columns.
            if len(cells) < 3 or cells[1]._tc is cells[2]._tc:
                prototype = row._tr
                break
    if prototype is None:
        return {}
    prototype_xml = copy.deepcopy(prototype)
    # The prototype may contain a legacy Tx_* bookmark. Cloning it once per row
    # would emit duplicate bookmark names and ids, which Word reports as a
    # damaged document, so the clone is stripped of bookmarks first.
    for tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
        for element in list(prototype_xml.iter(tag)):
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)

    body = table._tbl
    for row_element in list(body.tr_lst):
        body.remove(row_element)

    kp_anchors: dict[str, str] = {}
    bookmark_id = 9500
    for item in provisions:
        if not bool(item.get("include", True)):
            continue
        field = str(item.get("field", "")).strip()
        value = str(item.get("value", "")).strip()
        if not field and not value:
            continue
        body.append(copy.deepcopy(prototype_xml))
        row = table.rows[-1]
        label_paragraph = _set_cell_text(row.cells[0], field)
        paragraph = _set_cell_text(row.cells[1], value)

        # Anchor the row so cross-references in the body can jump back up to it.
        anchor = _kp_anchor_name(field)
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), anchor)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        properties = label_paragraph._p.find(qn("w:pPr"))
        label_paragraph._p.insert(1 if properties is not None else 0, start)
        label_paragraph._p.append(end)
        kp_anchors[field] = anchor
        bookmark_id += 1

        if bool(item.get("link")):
            section_anchor = anchors.get(str(item.get("section", "")).strip())
            if section_anchor:
                _append_hyperlink(paragraph, section_anchor, f"  (see Section {item['section']})")
    return kp_anchors


def _section_contains_ref(paragraphs: list[Paragraph], bookmark_name: str) -> bool:
    pattern = re.compile(rf"\bREF\s+{re.escape(bookmark_name)}\b", re.IGNORECASE)
    for paragraph in paragraphs:
        for instruction in paragraph._p.iter(qn("w:instrText")):
            if pattern.search(instruction.text or ""):
                return True
    return False


def _remove_empty_key_provision_rows(document: Document) -> None:
    """Collapse blank KPS rows so the visible table matches selected provisions."""
    if not document.tables:
        return
    table = document.tables[0]
    for row in reversed(table.rows):
        if all(not cell.text.strip() for cell in row.cells):
            parent = row._tr.getparent()
            if parent is not None:
                parent.remove(row._tr)


def _clean_legacy_rent_placeholders(document: Document) -> None:
    """Remove old Base Rent / Option table placeholder text."""
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "Table End" in paragraph.text or "DELETE TABLE" in paragraph.text:
                        paragraph.text = paragraph.text.replace("DELETE TABLE", "").replace("Table End", "").strip()


def _insert_schedule_table(document: Document, anchor: Paragraph, title: str, rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    heading = _insert_after(anchor, title, "")
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, header in zip(table.rows[0].cells, ["Period", "Monthly Rent", "Annual Rent"]):
        cell.text = header
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True
    for row_data in rows:
        cells = table.add_row().cells
        cells[0].text = str(row_data.get("Period", ""))
        cells[1].text = str(row_data.get("Monthly Rent", ""))
        cells[2].text = str(row_data.get("Annual Rent", ""))
    # The table is initially appended to the document; relocate it beneath the heading.
    heading._p.addnext(table._tbl)


def _insert_rent_schedules(document: Document, rent_schedules: dict[str, Any] | None) -> None:
    if not rent_schedules:
        return
    base_rows = list(rent_schedules.get("base", []))
    option_rows = list(rent_schedules.get("options", []))
    if not base_rows and not option_rows:
        return
    _replace_bookmark_text(document, "Tx_BaseRent", "")
    _clean_legacy_rent_placeholders(document)
    sections = {section["number"]: section for section in scan_sections(document)}
    paragraphs = list(document.paragraphs)
    base_section = sections.get("5")
    option_section = sections.get("5.3") or base_section
    if base_section and base_rows:
        _insert_schedule_table(document, paragraphs[base_section["start"]], "Base Rent Table", base_rows)
    if option_section and option_rows:
        anchor_index = max(option_section["start"], option_section["end"] - 1)
        _insert_schedule_table(document, paragraphs[anchor_index], "Option Rent Table", option_rows)


def _insert_after(paragraph: Paragraph, title: str, text: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if paragraph.style:
        new_paragraph.style = paragraph.style
    heading = new_paragraph.add_run(f"{title}.  ")
    heading.bold = True
    new_paragraph.add_run(text.strip())
    return new_paragraph


def _set_update_fields(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _clean_drafting_formatting(document: Document) -> None:
    """Accept revisions and remove template drafting markup from every Word story."""
    roots = []
    seen = set()

    # Main document plus headers/footers, comments and other OOXML story parts.
    for part in document.part.package.parts:
        root = getattr(part, "_element", None)
        if root is None:
            root = getattr(part, "element", None)
        if root is not None and id(root) not in seen:
            roots.append(root)
            seen.add(id(root))

    for root in roots:
        # Accept tracked insertions by unwrapping their contents into the document.
        for insertion in list(root.iter(qn("w:ins"))):
            parent = insertion.getparent()
            if parent is None:
                continue
            position = parent.index(insertion)
            for child in list(insertion):
                insertion.remove(child)
                parent.insert(position, child)
                position += 1
            parent.remove(insertion)

        # Reject deletions and remove comment anchors/references.
        remove_tags = (qn("w:del"), qn("w:commentRangeStart"), qn("w:commentRangeEnd"), qn("w:commentReference"))
        for tag in remove_tags:
            for element in list(root.iter(tag)):
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)

        # The source template uses red/blue/highlight/shading as drafting guidance.
        # A built lease is a clean document, so retain typography/layout but strip markup.
        for run_properties in root.iter(qn("w:rPr")):
            for tag in (qn("w:highlight"), qn("w:color"), qn("w:shd")):
                for element in list(run_properties.findall(tag)):
                    run_properties.remove(element)
        for tag in (qn("w:shd"), qn("w:highlight")):
            for element in list(root.iter(tag)):
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)

    settings = document.settings._element
    for tag in (qn("w:trackRevisions"), qn("w:showInsDel"), qn("w:showMarkup")):
        for element in list(settings.findall(tag)):
            settings.remove(element)


def _remove_template_markers(document: Document) -> None:
    """Remove Word-template bookmark labels such as (#Property) from output text."""
    marker_in_parentheses = re.compile(r"\s*\(\s*#[A-Za-z_][A-Za-z0-9_]*\s*\)")
    bare_marker = re.compile(r"\s+#[A-Za-z_][A-Za-z0-9_]*")
    roots = []
    seen = set()
    for part in document.part.package.parts:
        root = getattr(part, "_element", None)
        if root is None:
            root = getattr(part, "element", None)
        if root is not None and id(root) not in seen:
            roots.append(root)
            seen.add(id(root))
    for root in roots:
        # Hyperlinked display markers are split across separate runs: "(" + #Property + ")".
        for hyperlink in list(root.iter(qn("w:hyperlink"))):
            hyperlink_text = "".join(
                item.text or "" for item in hyperlink.iter(qn("w:t"))
            ).strip()
            if not re.fullmatch(r"#[A-Za-z_][A-Za-z0-9_]*", hyperlink_text):
                continue
            parent = hyperlink.getparent()
            if parent is None:
                continue
            index = parent.index(hyperlink)
            neighbors = []
            if index > 0:
                neighbors.append(parent[index - 1])
            if index + 1 < len(parent):
                neighbors.append(parent[index + 1])
            parent.remove(hyperlink)
            for neighbor in neighbors:
                neighbor_text = "".join(item.text or "" for item in neighbor.iter(qn("w:t"))).strip()
                if neighbor_text in {"(", ")"} and neighbor.getparent() is parent:
                    parent.remove(neighbor)
        for text_element in root.iter(qn("w:t")):
            original = text_element.text or ""
            cleaned = bare_marker.sub("", marker_in_parentheses.sub("", original))
            if re.fullmatch(r"#[A-Za-z_][A-Za-z0-9_]*", cleaned.strip()):
                cleaned = ""
            if cleaned != original:
                text_element.text = cleaned
                if cleaned.startswith(" ") or cleaned.endswith(" "):
                    text_element.set(qn("xml:space"), "preserve")


def _clean_drafting_notes(document: Document) -> None:
    """Remove template directions and scratch clauses before the formal lease body."""
    paragraphs = list(document.paragraphs)
    body_start = next(
        (i for i, paragraph in enumerate(paragraphs)
         if paragraph.text.strip().startswith("This LEASE AGREEMENT")),
        None,
    )
    if body_start is not None:
        for paragraph in reversed(paragraphs[2:body_start]):
            _delete_paragraph(paragraph)

    scratch_markers = ("DELETE ABOVE", "LEAVE FOR TEMPLATE", "Replace Sec 1 if full guarantee")
    for paragraph in list(document.paragraphs):
        if any(marker.lower() in paragraph.text.lower() for marker in scratch_markers):
            _delete_paragraph(paragraph)


def build_word_document(
    template_path: str | Path,
    section_choices: dict[str, dict[str, Any]],
    bookmark_values: dict[str, str] | None = None,
    included_bookmarks: set[str] | list[str] | None = None,
    linked_provisions: list[dict[str, Any]] | None = None,
    custom_provisions: list[dict[str, Any]] | None = None,
    rent_schedules: dict[str, Any] | None = None,
    additional_choices: list[dict[str, Any]] | None = None,
    clean_drafting_notes: bool = True,
    document_title: str = "MSP Lease Draft",
    preserve_template_structure: bool = False,
    key_provision_rows: list[dict[str, Any]] | None = None,
    token_report: dict[str, Any] | None = None,
) -> bytes:
    """Build a redline-ready DOCX and return its bytes.

    With preserve_template_structure the output is meant to be re-used as a base
    template rather than signed, so every key-provision row is kept even when its
    value is blank. Dropping those rows would delete the bookmarks with them and
    the published file would silently lose those provisions.
    """
    document = Document(str(template_path))

    # Convert legacy REF fields before anything else so section text, clause
    # replacements and the summary table all speak the same KP: language.
    if key_provision_rows is not None:
        convert_ref_fields_to_tokens(
            document,
            {
                str(row.get("bookmark", "")): str(row.get("field", ""))
                for row in key_provision_rows
                if row.get("bookmark") and row.get("field")
            },
        )

    sections = scan_sections(document)
    original_paragraphs = list(document.paragraphs)

    for section in reversed(sections):
        choice = section_choices.get(section["number"], {})
        clause_paragraphs = original_paragraphs[section["start"]:section["end"]]
        if not bool(choice.get("include", True)):
            for paragraph in reversed(clause_paragraphs):
                _delete_paragraph(paragraph)
            continue
        replacement = str(choice.get("replacement_text") or "").strip()
        if replacement:
            first = clause_paragraphs[0]
            title = str(choice.get("title") or section["title"]).strip()
            first.text = f"Section {section['number']}.  {title}.  {replacement}"
            for paragraph in reversed(clause_paragraphs[1:]):
                _delete_paragraph(paragraph)

    # The app owns the provision list: rebuild the summary table from it and drop
    # the legacy bookmark path entirely. Sections are already included/excluded
    # by this point, so anchors only exist for sections that survived.
    app_owned_provisions = key_provision_rows is not None
    kp_row_anchors: dict[str, str] = {}
    if app_owned_provisions:
        section_anchors = _add_section_anchors(document)
        kp_row_anchors = _rebuild_key_provision_table(document, key_provision_rows, section_anchors)

    if bookmark_values and not app_owned_provisions:
        values = dict(bookmark_values)
        if "Tx_Landlord" in values:
            values.setdefault("Tx_Landlord1", values["Tx_Landlord"])
        for bookmark_name, value in values.items():
            _replace_bookmark_text(document, bookmark_name, str(value))

    if included_bookmarks is not None and not preserve_template_structure and not app_owned_provisions:
        _apply_key_provision_inclusions(document, set(included_bookmarks))

    if custom_provisions and not app_owned_provisions and document.tables:
        table = document.tables[0]
        for item in custom_provisions:
            if not item.get("include"):
                continue
            field = str(item.get("field", "Key Provision")).strip()
            value = str(item.get("value", "")).strip()
            if not field or not value:
                continue
            row = table.add_row()
            row.cells[0].text = field
            row.cells[1].text = value
            row.cells[2].text = value

    _insert_rent_schedules(document, rent_schedules)

    current_sections = {section["number"]: section for section in scan_sections(document)}
    current_paragraphs = list(document.paragraphs)

    # Linking is a hyperlink in the summary table now, not text copied into the
    # clause, so the old injection is skipped once the app owns the provisions.
    if linked_provisions and not app_owned_provisions:
        grouped_links: dict[str, list[dict[str, Any]]] = {}
        for item in linked_provisions:
            section_number = str(item.get("section", ""))
            if item.get("include") and section_number and str(item.get("value", "")).strip():
                grouped_links.setdefault(section_number, []).append(item)
        for section_number, items in grouped_links.items():
            anchor_section = current_sections.get(section_number)
            if not anchor_section:
                continue
            section_paragraphs = current_paragraphs[anchor_section["start"]:anchor_section["end"]]
            anchor = section_paragraphs[-1]
            for item in items:
                bookmark_name = str(item.get("bookmark", ""))
                if bookmark_name and _section_contains_ref(section_paragraphs, bookmark_name):
                    continue
                anchor = _insert_after(anchor, str(item.get("field", "Key Provision")), str(item["value"]))

    if additional_choices:
        for item in additional_choices:
            if not item.get("include") or not str(item.get("text", "")).strip():
                continue
            anchor_section = current_sections.get(str(item.get("insert_after", "11")))
            if not anchor_section:
                continue
            anchor = current_paragraphs[anchor_section["end"] - 1]
            _insert_after(anchor, str(item.get("title", "Additional Provision")), str(item["text"]))

    _remove_template_markers(document)
    if not preserve_template_structure and not app_owned_provisions:
        _remove_empty_key_provision_rows(document)

    if clean_drafting_notes:
        _clean_drafting_notes(document)
        _clean_drafting_formatting(document)

    # Substitution runs last on purpose. _clean_drafting_formatting strips every
    # w:color to clear the template's red/blue drafting guidance, which would
    # take the green cross-reference styling with it. Running afterwards also
    # means tracked insertions are already resolved into the text.
    if app_owned_provisions:
        # An excluded provision resolves to nothing rather than leaving a broken
        # reference; the caller is told so the gap can be reviewed.
        excluded_names = {
            str(row.get("field", "")) for row in key_provision_rows
            if not bool(row.get("include", True))
        }
        values = {
            str(row.get("field", "")): (
                "" if str(row.get("field", "")) in excluded_names else str(row.get("value", ""))
            )
            for row in key_provision_rows if str(row.get("field", "")).strip()
        }
        referenced, unresolved = substitute_kp_tokens(document, values, kp_row_anchors)
        if token_report is not None:
            token_report["referenced"] = sorted(referenced)
            token_report["unresolved"] = sorted(unresolved)
            token_report["blank"] = sorted(referenced & excluded_names)

    document.core_properties.title = document_title
    document.core_properties.subject = "Generated by MSP Lease Builder"
    _set_update_fields(document)
    output = BytesIO()
    document.save(output)
    return output.getvalue()

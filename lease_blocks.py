"""Read the master lease's highlight convention into a menu of choices.

The master `.docx` is both a readable Word document and the machine-readable
menu the Lease Builder offers. Highlight colour carries the meaning:

    (none)      in the lease by default; still removable, just not surfaced
    green       lives in the master, OUT of the lease unless opted in
    cyan        one of an adjacent set — pick exactly one, or none
    yellow      an optional key provision (front matter)
    lightGray   scaffolding: drafting directions, stripped and never rendered
    red         scaffolding: notes to self, stripped and never rendered

Two structural rules do work that colour cannot:

**Indent groups a block.** A block begins at a highlighted paragraph and
absorbs the highlighted paragraphs that follow it at a deeper indent. So
"Conditional Reduction" carries "Application of Credit" and "Reinstatement
Rights" with it, and switching the parent off takes the children too.

**Separators divide the choices inside a cyan set.** A run of cyan is one
pick-one set, and each rule of asterisks inside it ends one option and starts
the next — so one separator means two choices, two separators mean three. An
option is everything between two separators, however many paragraphs that is.
Separators only ever appear inside cyan, and are stripped on the way out.

Containers are numbered sections *and* the exhibits after them. Without that
second half every exhibit collapses into Section 56, which is where the
Landlord's Work phases and the guaranty variants live — reading them as
alternatives to each other would let you pick Gas or Water but not both.

Pure data: takes a path, returns dictionaries. No Streamlit, no app state.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

# Meaning of each highlight, and whether it survives into a lease.
OPTIONAL_OUT = "green"      # present in the master, absent from the lease
PICK_ONE = "cyan"           # one of an adjacent set
KEY_PROVISION = "yellow"    # optional key provision, front matter
SCAFFOLDING = ("lightGray", "red")

SECTION_RE = re.compile(r"^\s*Section\s+(\d+(?:\.\d+)?)\s*\.?\s*(.*)$", re.IGNORECASE)
EXHIBIT_RE = re.compile(r'^\s*[“"\']?\s*(Exhibit\s+[A-Z0-9]+)\s*[”"\']?\s*$', re.IGNORECASE)
# Centred all-caps headings that open their own container.
NAMED_CONTAINER_RE = re.compile(
    r"^\s*(GUARANTY OF LEASE|SIGNATURES|TENANT FORMS\.?)\s*$", re.IGNORECASE
)
# A run-in heading: the bold lead-in that names a block.
RUN_IN_RE = re.compile(r"^\s*([^.:\n]{3,70}?)\s*[.:]\s")
# A rule of repeated symbols, used only inside a cyan set: it ends one choice
# and begins the next. N separators describe N+1 choices.
SEPARATOR_RE = re.compile(r"^[\s*_=~.\-\u2013\u2014]{3,}$")
# The colour legend at the top of the master describes the convention; it is
# not part of any lease.
LEGEND_RE = re.compile(r"^\s*(YELLOW|GREEN|CYAN|TEAL|RED|GREY|GRAY|MAGENTA)\s*=", re.IGNORECASE)


def _highlight(run: Any) -> str | None:
    properties = run._element.find(qn("w:rPr"))
    if properties is None:
        return None
    element = properties.find(qn("w:highlight"))
    return element.get(qn("w:val")) if element is not None else None


def _indent(paragraph: Any) -> int:
    """How deeply nested a paragraph is, as a level rather than a measurement.

    Word stores two things that look like indentation and disagree. `w:ind`
    is the visual left margin; `w:numPr/w:ilvl` is the list level, which is
    what the numbering (1, a, b, 2) actually reflects.

    In the guaranty they invert each other: "No Release" is item 1.d with
    ilvl=1 but w:ind=0, while "Good Guy Clause" is item 2 with ilvl=0 but
    w:ind=720. Reading the margin made Good Guy a child of No Release —
    the exact opposite of the document. The list level is authoritative
    wherever there is one.
    """
    properties = paragraph._element.find(qn("w:pPr"))
    if properties is None:
        return 0
    numbering = properties.find(qn("w:numPr"))
    if numbering is not None:
        level = numbering.find(qn("w:ilvl"))
        if level is not None:
            try:
                return int(level.get(qn("w:val")) or 0)
            except (TypeError, ValueError):
                return 0
    element = properties.find(qn("w:ind"))
    if element is None:
        return 0
    try:
        # Not in a list: fall back to the margin, expressed in the same units
        # as a list level so the two can be compared. Word's default step is
        # half an inch.
        return int(element.get(qn("w:left")) or 0) // 720
    except (TypeError, ValueError):
        return 0


def _alignment(paragraph: Any) -> str:
    properties = paragraph._element.find(qn("w:pPr"))
    if properties is None:
        return ""
    element = properties.find(qn("w:jc"))
    return (element.get(qn("w:val")) or "") if element is not None else ""


def paragraph_colour(paragraph: Any) -> str | None:
    """The one highlight that governs a paragraph, or None.

    A paragraph part-highlighted in a colour still counts as that colour: the
    convention works on whole paragraphs, and a half-marked one is far more
    likely to be a marking slip than a deliberate inline alternative. Those
    are reported rather than silently reinterpreted.
    """
    colours = {
        colour for run in paragraph.runs
        if (colour := _highlight(run)) and run.text.strip()
    }
    if not colours:
        return None
    for preferred in (PICK_ONE, OPTIONAL_OUT, KEY_PROVISION):
        if preferred in colours:
            return preferred
    return sorted(colours)[0]


def is_partly_highlighted(paragraph: Any) -> bool:
    """True when only some of a paragraph's text carries a highlight."""
    marked = "".join(run.text for run in paragraph.runs if _highlight(run)).strip()
    whole = paragraph.text.strip()
    return bool(marked) and len(marked) < len(whole) * 0.9


def block_name(text: str) -> str:
    """A block's label, taken from its run-in heading where there is one."""
    match = RUN_IN_RE.match(str(text or ""))
    if match:
        return match.group(1).strip()
    words = str(text or "").split()
    return " ".join(words[:7]) + ("…" if len(words) > 7 else "")


# ---------------------------------------------------------------------------
# Containers
# ---------------------------------------------------------------------------

def _container_for(paragraph: Any, text: str) -> tuple[str, str] | None:
    """(kind, label) when this paragraph opens a new container."""
    match = SECTION_RE.match(text)
    if match:
        return ("section", match.group(1))
    centred = _alignment(paragraph) == "center"
    match = EXHIBIT_RE.match(text)
    if match and centred:
        return ("exhibit", match.group(1).title())
    if centred and NAMED_CONTAINER_RE.match(text):
        return ("appendix", text.strip().title())
    return None


def _iter_body(document: Any):
    """Paragraphs and tables in the order they appear.

    python-docx's `document.paragraphs` skips table cells entirely, which is
    where the Key Provisions Summary lives — 45 of the yellow provisions were
    invisible until this walked the body itself.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "tbl", Table(child, document)


def _unique_cells(row: Any) -> list[Any]:
    """A row's cells with merges collapsed.

    python-docx returns a merged cell once per column it spans, so a two-column
    row built from a three-column table hands back the value cell twice and its
    paragraphs get counted twice with it.
    """
    seen, cells = set(), []
    for cell in row.cells:
        if id(cell._tc) not in seen:
            seen.add(id(cell._tc))
            cells.append(cell)
    return cells


def _key_provisions_from_table(table: Any) -> list[dict[str, Any]]:
    """Key provisions read off the summary table.

    The first cell names the provision, the second holds its value. The same
    convention that governs the body applies inside that cell: cyan paragraphs
    split by asterisk rules are competing versions of the value, yellow marks
    the whole provision optional, and grey or red is scaffolding.

    This is what replaced the Excel round trip. Alternatives used to exist only
    in a spreadsheet, which meant the master could not answer "what are my
    choices" on its own.
    """
    found = []
    for row in table.rows:
        cells = _unique_cells(row)
        if len(cells) < 2:
            continue
        field = cells[0].text.strip()
        paragraphs = [
            {"text": paragraph.text.strip(),
             "colour": paragraph_colour(paragraph),
             "separator": bool(SEPARATOR_RE.match(paragraph.text.strip()))}
            for paragraph in cells[1].paragraphs
            if paragraph.text.strip()
        ]
        optional = any(entry["colour"] == KEY_PROVISION for entry in paragraphs)
        # Drafting notes never reach a lease, and must not become a value.
        paragraphs = [e for e in paragraphs if e["colour"] not in SCAFFOLDING]
        if not field and not paragraphs:
            continue

        # Cyan runs split by separators are the competing values.
        alternatives, current, in_run = [], [], False
        base = []
        for entry in paragraphs:
            # Separator first: the rule is highlighted along with the choices
            # around it, so testing colour first would swallow it as content.
            if entry["separator"]:
                if in_run and current:
                    alternatives.append("\n".join(current))
                    current = []
                continue
            if entry["colour"] == PICK_ONE:
                in_run = True
                current.append(entry["text"])
            else:
                if in_run and current:
                    alternatives.append("\n".join(current))
                    current = []
                in_run = False
                base.append(entry["text"])
        if current:
            alternatives.append("\n".join(current))

        found.append({
            "name": block_name(field or (base[0] if base else "")),
            "field": field,
            "colour": KEY_PROVISION if optional else None,
            "optional": optional,
            "indent": 0,
            "text": "\n".join(base),
            "alternatives": alternatives,
            "children": [],
            "position": -1,
        })
    return found


def read_master(path: str | Path) -> dict[str, Any]:
    """Every container in the master, with its optional blocks and choices."""
    document = Document(str(path))

    containers: list[dict[str, Any]] = []
    current = {"kind": "front", "label": "Key Provisions Summary",
               "title": "", "paragraphs": []}
    warnings: list[str] = []
    pending_title = False

    for kind, item in _iter_body(document):
        if kind == "tbl":
            current.setdefault("table_blocks", []).extend(_key_provisions_from_table(item))
            continue
        paragraph = item
        text = paragraph.text.strip()
        if not text:
            continue
        if LEGEND_RE.match(text):
            continue  # the colour key, not lease language
        if SEPARATOR_RE.match(text):
            current["paragraphs"].append({"text": "", "colour": None,
                                          "indent": 0, "separator": True})
            continue

        opened = _container_for(paragraph, text)
        if opened:
            containers.append(current)
            kind, label = opened
            title = ""
            if kind == "section":
                # "Section 8.  Security Deposit.  Tenant will be subject to…"
                # The title is only the run-in heading; the rest of that
                # paragraph is the clause body and made every entry in the
                # index a wall of text.
                rest = SECTION_RE.match(text).group(2).strip()
                stop = re.search(r"[.:]\s", rest)
                title = (rest[:stop.start()] if stop and stop.start() <= 70
                         else rest[:70]).strip(" .\t")
            current = {"kind": kind, "label": label, "title": title, "paragraphs": []}
            pending_title = kind in ("exhibit", "appendix")
            continue

        # The line under an exhibit heading is its title, not body text.
        if pending_title and _alignment(paragraph) == "center" and len(text) < 80:
            current["title"] = text
            pending_title = False
            continue
        pending_title = False

        colour = paragraph_colour(paragraph)
        if colour in SCAFFOLDING:
            continue  # drafting directions never reach a lease
        if colour and is_partly_highlighted(paragraph):
            warnings.append(
                f"{current['label']}: “{text[:60]}…” is only part-highlighted "
                f"({colour}). The whole paragraph will be treated as {colour}."
            )
        current["paragraphs"].append(
            {"text": text, "colour": colour, "indent": _indent(paragraph)}
        )

    containers.append(current)
    for container in containers:
        container["blocks"] = (container.pop("table_blocks", [])
                               + _group_blocks(container["paragraphs"]))
        # The unhighlighted paragraphs are the bulk of the lease. Dropping them
        # left the document view showing only the parts with decisions in them,
        # which is the opposite of "what does this actually say".
        container["body"] = container.pop("paragraphs")

    return {"source": Path(path).name, "containers": containers, "warnings": warnings}


# ---------------------------------------------------------------------------
# Blocks and choice groups
# ---------------------------------------------------------------------------

def _group_blocks(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Highlighted paragraphs collapsed into blocks and pick-one sets.

    Cyan is handled first and differently. A run of cyan paragraphs is one
    pick-one set, and the rules of asterisks inside it delimit the options —
    so an option is *everything between two separators*, however many
    paragraphs that is. Exhibit C is the case that proves it: AS IS/WHERE IS
    is one option, and the entire Phase 1 and Phase 2 build-out is the other.

    Green and yellow use the indent rule instead: a block absorbs the deeper
    indented highlighted paragraphs beneath it.
    """
    blocks: list[dict[str, Any]] = []
    consumed: set[int] = set()
    group = 0

    index = 0
    while index < len(paragraphs):
        if index in consumed or paragraphs[index].get("separator") \
                or paragraphs[index]["colour"] != PICK_ONE:
            index += 1
            continue

        # Extend over every cyan paragraph and every separator between them.
        end = index
        last_cyan = index
        while end < len(paragraphs):
            entry = paragraphs[end]
            if entry.get("separator"):
                end += 1
                continue
            if entry["colour"] != PICK_ONE:
                break
            last_cyan = end
            end += 1
        run = paragraphs[index:last_cyan + 1]
        consumed.update(range(index, last_cyan + 1))

        # Split the run into options at the separators.
        options: list[list[dict[str, Any]]] = [[]]
        for entry in run:
            if entry.get("separator"):
                options.append([])
            else:
                options[-1].append(entry)
        options = [option for option in options if option]

        group += 1
        for option in options:
            head, *rest = option
            blocks.append({
                "name": block_name(head["text"]),
                "colour": PICK_ONE,
                "indent": head["indent"],
                "text": head["text"],
                "children": [entry["text"] for entry in rest],
                "position": paragraphs.index(head),
                # A lone cyan block is not a choice between anything, so only
                # a real set gets a group.
                **({"choice_group": group} if len(options) > 1 else {}),
            })
        index = last_cyan + 1

    # Green and yellow: the indent rule.
    index = 0
    while index < len(paragraphs):
        entry = paragraphs[index]
        if index in consumed or entry.get("separator") or not entry["colour"]:
            index += 1
            continue
        children: list[str] = []
        cursor = index + 1
        while cursor < len(paragraphs):
            following = paragraphs[cursor]
            if following.get("separator"):
                break
            if not following["colour"] or following["indent"] <= entry["indent"]:
                break
            children.append(following["text"])
            consumed.add(cursor)
            cursor += 1
        blocks.append({
            "name": block_name(entry["text"]),
            "colour": entry["colour"],
            "indent": entry["indent"],
            "text": entry["text"],
            "children": children,
            "position": index,
        })
        index = cursor

    blocks.sort(key=lambda block: block["position"])
    return blocks


# ---------------------------------------------------------------------------
# Report
# ---------------------------------------------------------------------------

def report(master: dict[str, Any]) -> str:
    """A human-readable reading of the master, to check before it is trusted."""
    lines = [f"Master: {master['source']}", ""]
    totals = {OPTIONAL_OUT: 0, PICK_ONE: 0, KEY_PROVISION: 0}

    for container in master["containers"]:
        blocks = container["blocks"]
        if not blocks:
            continue
        heading = container["label"]
        if container["kind"] == "section":
            heading = f"Section {container['label']}. {container['title']}"
        elif container["title"]:
            heading = f"{container['label']} — {container['title']}"
        lines.append(heading)

        groups: dict[int, list[dict]] = {}
        for block in blocks:
            totals[block["colour"]] = totals.get(block["colour"], 0) + 1
            if block.get("choice_group"):
                groups.setdefault(block["choice_group"], []).append(block)

        shown = set()
        for block in blocks:
            group = block.get("choice_group")
            if group:
                if group in shown:
                    continue
                shown.add(group)
                members = groups[group]
                lines.append(f"    PICK ONE of {len(members)}:")
                for member in members:
                    extra = f" (+{len(member['children'])})" if member["children"] else ""
                    lines.append(f"        · {member['name']}{extra}")
            else:
                mark = {OPTIONAL_OUT: "OFF by default",
                        PICK_ONE: "optional (single)",
                        KEY_PROVISION: "key provision"}.get(block["colour"], block["colour"])
                extra = f" (+{len(block['children'])} sub)" if block["children"] else ""
                lines.append(f"    [ ] {block['name']}{extra}  — {mark}")
        lines.append("")

    lines.append(f"TOTALS  off-by-default: {totals.get(OPTIONAL_OUT, 0)}   "
                 f"pick-one: {totals.get(PICK_ONE, 0)}   "
                 f"key provisions: {totals.get(KEY_PROVISION, 0)}")
    if master["warnings"]:
        lines.append("")
        lines.append("WARNINGS")
        for warning in master["warnings"]:
            lines.append(f"  ! {warning}")
    return "\n".join(lines)

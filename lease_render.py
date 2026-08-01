"""Word renderer: builds a lease from rules, with no base .docx.

This is where `lease_format.py` (how it should look) and `lease_markup.py`
(what it says) meet python-docx. Everything starts from an empty Document and
is styled from the settings dict, so there is no template file to lose, no
inherited formatting to fight, and no bookmarks to preserve by accident.

Scope of this module today (spec build order steps 3 and 5's run handling):
page setup, styles, footer, title block, and paragraph/run rendering including
cross-reference hyperlinks. The Key Provisions table, rent tables, signature
block and exhibits arrive in steps 4, 6, 7 and 8.
"""

from __future__ import annotations

import re
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import lease_format as lf
import lease_markup as lm

# Bookmark name prefixes. Word treats a duplicate bookmark ID as file damage, so
# every name is issued through _Bookmarks below rather than built ad hoc.
KP_ANCHOR_PREFIX = "_MSP_KP_"
SECTION_ANCHOR_PREFIX = "_MSP_Sec_"

_ALIGNMENTS = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _element(tag: str, **attributes: str) -> Any:
    node = OxmlElement(tag)
    for name, value in attributes.items():
        node.set(qn(name.replace("__", ":")), value)
    return node


class _Bookmarks:
    """Issues unique bookmark ids and refuses to reuse a name.

    Word reports duplicate bookmark IDs as a damaged file, and the failure is
    silent until someone opens the lease, so uniqueness is enforced here rather
    than hoped for.
    """

    def __init__(self) -> None:
        self._next_id = 1
        self._used: set[str] = set()

    def sanitize(self, raw: str, prefix: str) -> str:
        """Word bookmark names: letters, digits and underscores, <= 40 chars."""
        stem = re.sub(r"[^A-Za-z0-9]+", "_", str(raw)).strip("_")
        name = (prefix + stem)[:40].rstrip("_")
        if not name or name == prefix.rstrip("_"):
            name = f"{prefix}Unnamed"
        candidate = name
        suffix = 2
        while candidate in self._used:
            tail = f"_{suffix}"
            candidate = name[: 40 - len(tail)] + tail
            suffix += 1
        return candidate

    def wrap(self, paragraph: Any, name: str) -> str:
        """Bookmark the whole paragraph. Returns the name actually used."""
        if name in self._used:
            name = self.sanitize(name, "")
        self._used.add(name)
        identifier = str(self._next_id)
        self._next_id += 1
        start = _element("w:bookmarkStart", w__id=identifier, w__name=name)
        end = _element("w:bookmarkEnd", w__id=identifier)
        paragraph._p.insert(0, start)
        paragraph._p.append(end)
        return name

    @property
    def names(self) -> set[str]:
        return set(self._used)


def _add_field(paragraph: Any, instruction: str, placeholder: str = "1") -> None:
    """A Word field such as PAGE or NUMPAGES."""
    run = paragraph.add_run()
    run._r.append(_element("w:fldChar", w__fldCharType="begin"))
    instr = _element("w:instrText", xml__space="preserve")
    instr.text = f" {instruction} "
    run._r.append(instr)
    run._r.append(_element("w:fldChar", w__fldCharType="separate"))
    text = _element("w:t")
    text.text = placeholder
    run._r.append(text)
    run._r.append(_element("w:fldChar", w__fldCharType="end"))


def _add_internal_hyperlink(paragraph: Any, anchor: str, text: str,
                            settings: dict[str, Any]) -> Any:
    """Green underlined jump to a bookmark, used for [KP:Name] references."""
    link = _element("w:hyperlink", w__anchor=anchor)
    run = _element("w:r")
    properties = _element("w:rPr")
    color = _element("w:color", w__val=lf.normalize_hex_color(settings.get("kp_link_color")))
    properties.append(color)
    if settings.get("kp_link_underline", True):
        properties.append(_element("w:u", w__val="single"))
    run.append(properties)
    node = _element("w:t", xml__space="preserve")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)
    return link


def bookmark_names(document: Any) -> list[str]:
    return [b.get(qn("w:name")) for b in document.element.body.iter(qn("w:bookmarkStart"))]


def dangling_anchors(document: Any) -> list[str]:
    """Hyperlink targets with no matching bookmark.

    A dangling internal link does not corrupt the file — it simply does not
    jump — so it has to be checked for deliberately. Until step 4 builds the
    Key Provisions table, every provision anchor is expected to appear here.
    """
    defined = set(bookmark_names(document))
    anchors = {
        h.get(qn("w:anchor"))
        for h in document.element.body.iter(qn("w:hyperlink"))
        if h.get(qn("w:anchor"))
    }
    return sorted(anchor for anchor in anchors if anchor not in defined)


def set_update_fields_on_open(document: Any) -> None:
    """Ask Word to refresh PAGE fields when the document opens."""
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is None:
        settings.append(_element("w:updateFields", w__val="true"))


# ---------------------------------------------------------------------------
# Document skeleton
# ---------------------------------------------------------------------------

def apply_page_setup(document: Any, settings: dict[str, Any]) -> None:
    width, height = lf.page_dimensions(settings)
    for section in document.sections:
        section.page_width = Inches(width)
        section.page_height = Inches(height)
        section.top_margin = Inches(settings["margin_top_in"])
        section.bottom_margin = Inches(settings["margin_bottom_in"])
        section.left_margin = Inches(settings["margin_left_in"])
        section.right_margin = Inches(settings["margin_right_in"])
        section.footer_distance = Inches(0.5)


def build_styles(document: Any, settings: dict[str, Any]) -> None:
    """Set Normal, and add the named styles the renderer uses.

    Everything hangs off Normal so a font change in the form moves the whole
    document rather than the body only.
    """
    normal = document.styles["Normal"]
    normal.font.name = settings["body_font"]
    normal.font.size = Pt(settings["body_size_pt"])
    # East-Asian and complex-script names must be set too or Word substitutes.
    rpr = normal.element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = _element("w:rFonts")
        rpr.insert(0, fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        fonts.set(qn(attribute), settings["body_font"])

    paragraph_format = normal.paragraph_format
    paragraph_format.alignment = _ALIGNMENTS.get(settings["body_alignment"], WD_ALIGN_PARAGRAPH.JUSTIFY)
    paragraph_format.line_spacing = settings["body_line_spacing"]
    paragraph_format.space_after = Pt(settings["space_after_pt"])
    paragraph_format.space_before = Pt(0)
    paragraph_format.first_line_indent = Inches(settings["first_line_indent_in"])
    # Widow/orphan control: a clause that breaks with one line stranded reads as
    # a drafting error even though it is only pagination.
    paragraph_format.widow_control = True


def add_footer(document: Any, settings: dict[str, Any]) -> None:
    """Counsel document ID at the left, page number at the right."""
    position = str(settings.get("page_number_position", "right"))
    doc_id = str(settings.get("footer_doc_id", "")).strip()
    content_width = lf.content_width_in(settings)

    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for run in list(paragraph.runs):
            run._r.getparent().remove(run._r)
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Inches(0)

        tab_stops = paragraph.paragraph_format.tab_stops
        if position == "center":
            tab_stops.add_tab_stop(Inches(content_width / 2), WD_TAB_ALIGNMENT.CENTER)
        elif position in ("right", "left"):
            tab_stops.add_tab_stop(Inches(content_width), WD_TAB_ALIGNMENT.RIGHT)

        if position == "left" and doc_id:
            # Page number first, then the id pushed to the right.
            _add_field(paragraph, "PAGE")
            paragraph.add_run("\t" + doc_id)
        else:
            if doc_id:
                paragraph.add_run(doc_id)
            if position != "none":
                paragraph.add_run("\t")
                _add_field(paragraph, "PAGE")

        for run in paragraph.runs:
            run.font.size = Pt(settings["footer_font_size_pt"])
            run.font.name = settings["body_font"]


def add_title_block(document: Any, settings: dict[str, Any]) -> list[Any]:
    """Centered, bold, underlined document title and summary title."""
    paragraphs = []
    for text in (settings.get("title_text", ""), settings.get("key_provisions_title", "")):
        if not str(text).strip():
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.space_after = Pt(settings["space_after_pt"])
        run = paragraph.add_run(str(text))
        run.bold = bool(settings.get("title_bold", True))
        run.underline = bool(settings.get("title_underline", True))
        run.font.size = Pt(settings["title_size_pt"])
        paragraphs.append(paragraph)
    return paragraphs


def new_document(settings: dict[str, Any] | None = None) -> tuple[Any, dict[str, Any]]:
    """An empty, fully configured document plus the normalized settings."""
    resolved = lf.normalize_settings(settings)
    document = Document()
    # A fresh Document starts with one empty paragraph; drop it so the title is
    # the first thing on page 1.
    for paragraph in list(document.paragraphs):
        paragraph._p.getparent().remove(paragraph._p)
    apply_page_setup(document, resolved)
    build_styles(document, resolved)
    add_footer(document, resolved)
    set_update_fields_on_open(document)
    return document, resolved


# ---------------------------------------------------------------------------
# Blocks and runs
# ---------------------------------------------------------------------------

class BlockRenderer:
    """Turns lease_markup Blocks into Word paragraphs.

    Holds the bookmark registry so cross-reference anchors stay unique across
    the whole document, and records which provisions were referenced but never
    defined — those must be visible, never silently blank.
    """

    def __init__(self, document: Any, settings: dict[str, Any],
                 values: dict[str, str] | None = None,
                 bookmarks: _Bookmarks | None = None) -> None:
        self.document = document
        self.settings = settings
        self.values = {str(k).strip().lower(): str(v) for k, v in (values or {}).items()}
        self.anchors: dict[str, str] = {}
        self.bookmarks = bookmarks or _Bookmarks()
        self.unresolved: set[str] = set()

    # -- anchors ---------------------------------------------------------
    def register_provision(self, field: str) -> str:
        """Reserve the bookmark a [KP:Name] reference will jump to."""
        key = str(field).strip().lower()
        if key not in self.anchors:
            self.anchors[key] = self.bookmarks.sanitize(field, KP_ANCHOR_PREFIX)
        return self.anchors[key]

    # -- paragraphs ------------------------------------------------------
    def _new_paragraph(self, block: lm.Block) -> Any:
        paragraph = self.document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.alignment = _ALIGNMENTS.get(self.settings["body_alignment"], WD_ALIGN_PARAGRAPH.JUSTIFY)

        if block.is_list:
            step = float(self.settings["subclause_level1_indent_in"])
            hanging = float(self.settings["subclause_level1_hanging_in"])
            left = step * (block.level + 1)
            fmt.left_indent = Inches(left)
            fmt.first_line_indent = Inches(-hanging)
        else:
            fmt.left_indent = Inches(float(self.settings["subclause_level1_indent_in"]) * block.level)
            fmt.first_line_indent = Inches(self.settings["first_line_indent_in"])
        return paragraph

    def _add_text_run(self, paragraph: Any, run: lm.Run) -> Any:
        word_run = paragraph.add_run(run.text)
        word_run.bold = run.bold
        word_run.italic = run.italic
        return word_run

    def render_block(self, block: lm.Block) -> Any | None:
        """One block. Returns the paragraph, or None for a page break."""
        if block.kind == lm.BLOCK_PAGE_BREAK:
            paragraph = self.document.add_paragraph()
            paragraph.add_run().add_break(6)  # WD_BREAK.PAGE
            return None
        if block.is_structural:
            # Rent tables and exhibits are steps 6 and 8; leave a marker
            # paragraph so nothing silently disappears in the meantime.
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(lm.to_markup([block]))
            run.italic = True
            return paragraph

        paragraph = self._new_paragraph(block)
        if block.number:
            marker = paragraph.add_run(block.number + "\t")
            marker.bold = bool(self.settings.get("subclause_lead_in_bold", True)) and (
                block.kind == lm.BLOCK_LETTERED
            )

        for run in block.runs:
            if run.is_break:
                paragraph.add_run().add_break()
            elif run.is_ref:
                self._add_reference(paragraph, run)
            else:
                self._add_text_run(paragraph, run)
        return paragraph

    def _add_reference(self, paragraph: Any, run: lm.Run) -> None:
        key = run.name.strip().lower()
        value = self.values.get(key)
        if value is None:
            # An unknown provision stays visible as its token. A blank here
            # would be a hole in a signed lease that nobody notices.
            self.unresolved.add(run.name)
            fallback = paragraph.add_run(f"[KP:{run.name}]")
            fallback.bold = run.bold
            fallback.italic = run.italic
            fallback.font.color.rgb = RGBColor(0xB0, 0x00, 0x20)
            return
        anchor = self.register_provision(run.name)
        for position, line in enumerate(str(value).split(lm.LINE_BREAK)):
            if position:
                paragraph.add_run().add_break()
            _add_internal_hyperlink(paragraph, anchor, line, self.settings)

    def render(self, blocks: Iterable[lm.Block]) -> list[Any]:
        return [p for p in (self.render_block(block) for block in blocks) if p is not None]


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def render_section(renderer: BlockRenderer, number: str, title: str, body: str) -> Any:
    """A section: run-in bold heading, body continuing on the same line.

    Decimal numbers are peers, not children, so they render identically.
    """
    settings = renderer.settings
    blocks = lm.parse_blocks(body)

    heading = renderer.document.add_paragraph()
    fmt = heading.paragraph_format
    fmt.alignment = _ALIGNMENTS.get(settings["body_alignment"], WD_ALIGN_PARAGRAPH.JUSTIFY)
    fmt.space_before = Pt(settings["section_space_before_pt"])
    fmt.first_line_indent = Inches(settings["section_first_line_indent_in"])
    fmt.keep_with_next = True
    fmt.tab_stops.add_tab_stop(Inches(settings["section_tab_stop_in"]))

    bold = bool(settings.get("section_heading_bold", True))
    label = heading.add_run(f"{settings['section_word']} {number}.")
    label.bold = bold
    heading.add_run("\t")
    if title:
        title_run = heading.add_run(f"{title}.")
        title_run.bold = bold
        heading.add_run("  ")

    renderer.bookmarks.wrap(heading, renderer.bookmarks.sanitize(number, SECTION_ANCHOR_PREFIX))

    # The first block continues on the heading line; the rest are their own
    # paragraphs. That is what "run-in" means and it is how the executed lease
    # reads.
    if blocks and not blocks[0].is_structural and not blocks[0].is_list:
        for run in blocks[0].runs:
            if run.is_break:
                heading.add_run().add_break()
            elif run.is_ref:
                renderer._add_reference(heading, run)
            else:
                renderer._add_text_run(heading, run)
        blocks = blocks[1:]

    renderer.render(blocks)
    return heading


def render_lease(sections: list[dict[str, Any]],
                 settings: dict[str, Any] | None = None,
                 values: dict[str, str] | None = None) -> tuple[Any, BlockRenderer]:
    """Title block plus every section. Tables and exhibits come in later steps."""
    document, resolved = new_document(settings)
    renderer = BlockRenderer(document, resolved, values=values)
    add_title_block(document, resolved)
    for section in sections:
        render_section(
            renderer,
            str(section.get("number", "")),
            str(section.get("title", "")),
            str(section.get("body", "")),
        )
    return document, renderer
